import requests
from bs4 import BeautifulSoup
from docx import Document

BASE_URL = "https://www.zodiacsign.com/zodiac-signs/{}"
SIGNS = ["aquarius", "pisces", "aries", "taurus", "gemini", "cancer","leo", "virgo", "libra", "scorpio", "sagittarius", "capricorn"]

def scrape_sign(sign):
    url = BASE_URL.format(sign + "/")
    resp = requests.get(url)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "html.parser")

    title = soup.find("h1").get_text(strip=True) if soup.find("h1") else sign.capitalize()

    attributes = {}
    for strong in soup.find_all("strong"):
        text = strong.get_text(strip=True)
        if ":" in text:
            key = text.split(":")[0].strip()
            value = strong.next_sibling.strip() if strong.next_sibling else ""
            attributes[key] = value

    sections = {}
    for header in soup.find_all(["h2", "h3"]):
        section_title = header.get_text(strip=True)
        content_parts = []
        for sibling in header.find_next_siblings():
            if sibling.name in ["h2", "h3"]:
                break
            if sibling.name == "p":
                content_parts.append(sibling.get_text(strip=True))
        if content_parts:
            sections[section_title] = "\n\n".join(content_parts)

    return title, attributes, sections


def save_to_word(data, filename="zodiac_info.docx"):
    doc = Document()
    doc.add_heading("Zodiac Signs Detailed Information", 0)

    for title, attributes, sections in data:
        doc.add_heading(title, level=1)

        
        if attributes:
            doc.add_heading("Attributes", level=2)
            for key, val in attributes.items():
                doc.add_paragraph(f"{key}: {val}")

        
        for section_title, content in sections.items():
            doc.add_heading(section_title, level=2)
            doc.add_paragraph(content)

        doc.add_page_break()

    doc.save(filename)
    print(f"Saved to {filename}")


if __name__ == "__main__":
    all_data = []
    for sign in SIGNS:
        print(f"Scraping {sign}...")
        try:
            all_data.append(scrape_sign(sign))
        except Exception as e:
            print(f"Failed to scrape {sign}: {e}")

    save_to_word(all_data)
